"""
File Analysis Service with LLM
تحلیل فایل‌های ضمیمه شده با استفاده از LLM
پشتیبانی از استخراج متن با OCR برای فایل‌های PDF و DOCX
"""

from typing import Optional, Dict, Any, List, Tuple
import base64
import io
import structlog
import requests

from app.llm.base import LLMConfig, LLMProvider, Message
from app.llm.openai_provider import OpenAIProvider
from app.config.settings import settings

logger = structlog.get_logger()

# ============================================================================
# File Type Detection
# ============================================================================

# Image extensions (supported formats)
IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.webp'}

# Text file extensions (can be extracted locally)
TEXT_FILE_EXTENSIONS = {'.pdf', '.txt', '.doc', '.docx', '.html', '.htm'}


def get_file_extension(url_or_filename: str) -> str:
    """Extract file extension from URL or filename"""
    import os
    from urllib.parse import urlparse, unquote
    
    # Parse URL and get path
    parsed = urlparse(url_or_filename)
    path = unquote(parsed.path)
    
    # Get extension
    _, ext = os.path.splitext(path)
    return ext.lower()


def is_image_file(url_or_filename: str) -> bool:
    """Check if file is an image"""
    ext = get_file_extension(url_or_filename)
    return ext in IMAGE_EXTENSIONS


def is_text_file(url_or_filename: str) -> bool:
    """Check if file is a text-based file that can be extracted locally"""
    ext = get_file_extension(url_or_filename)
    return ext in TEXT_FILE_EXTENSIONS


# ============================================================================
# Local Text Extraction with OCR Support
# ============================================================================

# Try importing optional libraries
try:
    from PyPDF2 import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

try:
    import fitz  # PyMuPDF for PDF images
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False


def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF content with OCR support for images"""
    text_parts = []
    
    # Method 1: Try PyMuPDF (better for images)
    if HAS_PYMUPDF:
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            for page_num, page in enumerate(doc):
                # Extract text
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(page_text)
                
                # Extract images and OCR them
                if HAS_OCR:
                    image_list = page.get_images()
                    for img_index, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            base_image = doc.extract_image(xref)
                            image_bytes = base_image["image"]
                            
                            # OCR the image
                            pil_image = Image.open(io.BytesIO(image_bytes))
                            ocr_text = pytesseract.image_to_string(pil_image, lang='fas+eng')
                            if ocr_text.strip():
                                text_parts.append(f"\n[متن استخراج شده از تصویر صفحه {page_num + 1}]:\n{ocr_text}")
                        except:
                            continue
            doc.close()
            
            if text_parts:
                return "\n".join(text_parts).strip()
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {e}")
    
    # Method 2: Fall back to PyPDF2
    if HAS_PDF:
        try:
            reader = PdfReader(io.BytesIO(content))
            for page in reader.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(page_text)
            
            if text_parts:
                return "\n".join(text_parts).strip()
        except Exception as e:
            return f"[خطا در خواندن PDF: {e}]"
    
    if not HAS_PDF and not HAS_PYMUPDF:
        return "[خطا: کتابخانه PyPDF2 یا PyMuPDF نصب نیست]"
    
    return "[فایل PDF بدون متن قابل استخراج]"


def extract_text_from_txt(content: bytes) -> str:
    """Extract text from TXT content"""
    try:
        # Try different encodings
        for encoding in ['utf-8', 'utf-16', 'cp1256', 'iso-8859-1']:
            try:
                return content.decode(encoding)
            except:
                continue
        return content.decode('utf-8', errors='ignore')
    except Exception as e:
        return f"[خطا در خواندن TXT: {e}]"


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX content with OCR support for images"""
    if not HAS_DOCX:
        return "[خطا: کتابخانه python-docx نصب نیست]"
    try:
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    text_parts.append(row_text)
        
        # Extract images and OCR them
        if HAS_OCR:
            import zipfile
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                for file_name in zf.namelist():
                    if file_name.startswith('word/media/') and file_name.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                        try:
                            image_data = zf.read(file_name)
                            pil_image = Image.open(io.BytesIO(image_data))
                            ocr_text = pytesseract.image_to_string(pil_image, lang='fas+eng')
                            if ocr_text.strip():
                                text_parts.append(f"\n[متن استخراج شده از تصویر]:\n{ocr_text}")
                        except:
                            continue
        
        return "\n".join(text_parts).strip() if text_parts else "[فایل DOCX بدون متن]"
    except Exception as e:
        return f"[خطا در خواندن DOCX: {e}]"


def extract_text_from_html(content: bytes) -> str:
    """Extract text from HTML content"""
    if not HAS_BS4:
        return "[خطا: کتابخانه beautifulsoup4 نصب نیست]"
    try:
        soup = BeautifulSoup(content, 'html.parser')
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        text = soup.get_text(separator='\n', strip=True)
        return text if text else "[فایل HTML بدون متن]"
    except Exception as e:
        return f"[خطا در خواندن HTML: {e}]"


def extract_text_from_file(content: bytes, extension: str) -> Tuple[str, str]:
    """
    Extract text from file based on extension.
    Returns: (text_content, error_message)
    """
    try:
        if extension == '.pdf':
            text = extract_text_from_pdf(content)
        elif extension == '.txt' or extension == '.md' or extension == '.csv' or extension == '.json' or extension == '.xml':
            text = extract_text_from_txt(content)
        elif extension in ['.doc', '.docx']:
            text = extract_text_from_docx(content)
        elif extension in ['.html', '.htm']:
            text = extract_text_from_html(content)
        else:
            return "", f"فرمت {extension} پشتیبانی نمی‌شود"
        
        # Limit text length to avoid token overflow
        max_chars = 15000
        if len(text) > max_chars:
            text = text[:max_chars] + "\n\n... [متن ادامه دارد - برش خورده]"
        
        return text, ""
        
    except Exception as e:
        return "", f"خطا در پردازش فایل: {e}"


async def download_file_content(url: str) -> bytes:
    """Download file content from URL"""
    import asyncio
    
    def _download():
        response = requests.get(url, timeout=30)
        response.raise_for_status()
        return response.content
    
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _download)


class FileAnalysisService:
    """سرویس تحلیل فایل با LLM"""
    
    def __init__(self):
        """Initialize file analysis service with LLM"""
        self.llm_config = LLMConfig(
            provider=LLMProvider.OPENAI_COMPATIBLE,
            model=settings.llm_model,
            api_key=settings.llm_api_key,
            base_url=settings.llm_base_url,
            temperature=0.2,
            max_tokens=2000,
        )
        self.llm = OpenAIProvider(self.llm_config)
        logger.info("FileAnalysisService initialized")
    
    async def analyze_file_from_url(
        self,
        file_url: str,
        user_query: str = ""
    ) -> Dict[str, Any]:
        """
        اگر فایل تصویر باشد → مستقیم به مدل ارسال می‌شود (input_image)
        اگر فایل متنی باشد → ابتدا متن استخراج شده و سپس به مدل ارسال می‌شود
        
        Args:
            file_url: آدرس فایل
            user_query: سوال کاربر (اختیاری)
            
        Returns:
            dict با کلیدهای answer, input_tokens, output_tokens یا error
        """
        from app.config.prompts import FileAnalysisPrompts
        
        ANALYSIS_PROMPT = FileAnalysisPrompts.get_analysis_prompt()
        ANALYSIS_USER_TEXT = FileAnalysisPrompts.get_analysis_user_text()
        MAX_TOKENS_RESPONSE = 8192
        
        try:
            # تشخیص نوع فایل
            if is_image_file(file_url):
                # Image file - use input_image (مانند فایل تست)
                logger.info(f"Image file detected, sending to vision model: {file_url}")
                
                input_content = [
                    {
                        "role": "user",
                        "content": [
                            {"type": "input_text", "text": f"{ANALYSIS_PROMPT}\n\n---\n\n{ANALYSIS_USER_TEXT}"},
                            {"type": "input_image", "image_url": file_url}
                        ]
                    }
                ]
                
                response = await self.llm.generate_responses_api(
                    messages=[],  # خالی چون از input_content استفاده می‌کنیم
                    reasoning_effort="high",
                    input_content=input_content,
                    max_tokens=MAX_TOKENS_RESPONSE
                )
                
                logger.info(f"Image analyzed", analysis_length=len(response.content))
                return {
                    "answer": response.content.strip(),
                    "input_tokens": response.usage.get("prompt_tokens", 0),
                    "output_tokens": response.usage.get("completion_tokens", 0)
                }
            
            elif is_text_file(file_url):
                # Text file - extract content locally with OCR support (مانند فایل تست)
                logger.info(f"Text file detected, extracting text locally: {file_url}")
                
                # دانلود فایل
                file_content = await download_file_content(file_url)
                logger.info(f"File downloaded: {len(file_content)} bytes")
                
                # استخراج متن
                extension = get_file_extension(file_url)
                extracted_text, error = extract_text_from_file(file_content, extension)
                
                if error:
                    logger.warning(f"Text extraction error: {error}")
                    return {"error": f"خطا در استخراج محتوا: {error}"}
                
                logger.info(f"Text extracted: {len(extracted_text)} characters")
                
                # ارسال متن به مدل
                input_content = f"{ANALYSIS_PROMPT}\n\n---\n\n{ANALYSIS_USER_TEXT}\n\n**محتوای فایل:**\n{extracted_text}"
                
                response = await self.llm.generate_responses_api(
                    messages=[],
                    reasoning_effort="high",
                    input_content=input_content,
                    max_tokens=MAX_TOKENS_RESPONSE
                )
                
                logger.info(f"Text file analyzed", analysis_length=len(response.content))
                return {
                    "answer": response.content.strip(),
                    "input_tokens": response.usage.get("prompt_tokens", 0),
                    "output_tokens": response.usage.get("completion_tokens", 0)
                }
            
            else:
                return {"error": "فرمت فایل پشتیبانی نمی‌شود"}
        
        except Exception as e:
            import traceback
            logger.error(f"Failed to analyze file: {e}")
            return {"error": str(e), "traceback": traceback.format_exc()}
    
    async def analyze_file(
        self,
        file_url: str,
        filename: str = "",
        file_type: str = "",
        user_query: str = "",
        language: str = "fa"
    ) -> str:
        """
        تحلیل یک فایل منفرد با LLM (متد قدیمی برای سازگاری)
        
        Args:
            file_url: آدرس فایل
            filename: نام فایل (اختیاری)
            file_type: نوع فایل (اختیاری)
            user_query: سوال کاربر (اختیاری)
            language: زبان
            
        Returns:
            تحلیل فایل (متن)
        """
        # استفاده از متد جدید
        result = await self.analyze_file_from_url(file_url, user_query)
        
        # برگرداندن فقط متن پاسخ برای سازگاری
        if "error" in result:
            return f"خطا در تحلیل فایل: {result['error']}"
        return result.get("answer", "")
    
    async def analyze_files(
        self,
        files_content: List[Dict[str, Any]],
        user_query: str,
        language: str = "fa"
    ) -> str:
        """
        تحلیل فایل‌های ضمیمه شده با LLM (متد قدیمی - فقط متن)
        
        Args:
            files_content: لیست فایل‌ها شامل:
                - filename: نام فایل
                - content: محتوای استخراج شده (متن)
                - file_type: نوع فایل
                - is_image: آیا تصویر است؟
            user_query: سوال کاربر
            language: زبان
            
        Returns:
            تحلیل کامل فایل‌ها
        """
        # استفاده از متد جدید با پشتیبانی از تصاویر
        return await self.analyze_files_with_images(files_content, user_query, language)
    
    async def analyze_files_with_images(
        self,
        files_content: List[Dict[str, Any]],
        user_query: str,
        language: str = "fa"
    ) -> str:
        """
        تحلیل فایل‌های ضمیمه شده با LLM (با پشتیبانی از تصاویر)
        
        این متد از Responses API با input_image برای تصاویر استفاده می‌کند.
        
        Args:
            files_content: لیست فایل‌ها شامل:
                - filename: نام فایل
                - content: محتوای استخراج شده (متن) - برای فایل‌های متنی
                - file_type: نوع فایل
                - is_image: آیا تصویر است؟
                - image_url: URL معتبر تصویر (presigned URL از MinIO)
            user_query: سوال کاربر
            language: زبان
            
        Returns:
            تحلیل کامل فایل‌ها
        """
        try:
            # جداسازی تصاویر و فایل‌های متنی
            images = [f for f in files_content if f.get('is_image') and f.get('image_url')]
            text_files = [f for f in files_content if not f.get('is_image')]
            
            # ساخت محتوای پیام
            content_parts = []
            
            # اضافه کردن prompt متنی
            text_prompt = self._build_analysis_prompt_for_mixed(
                text_files, images, user_query, language
            )
            content_parts.append({
                "type": "input_text",
                "text": text_prompt
            })
            
            # اضافه کردن تصاویر با input_image
            for img in images:
                content_parts.append({
                    "type": "input_image",
                    "image_url": img['image_url']
                })
            
            # ساخت input برای Responses API
            input_content = [
                {
                    "role": "user",
                    "content": content_parts
                }
            ]
            
            logger.info(
                "Sending files to LLM with images",
                text_files_count=len(text_files),
                images_count=len(images)
            )
            
            # ارسال به LLM با Responses API
            response = await self.llm.generate_responses_api(
                messages=[],  # خالی چون از input_content استفاده می‌کنیم
                reasoning_effort="medium",
                input_content=input_content,
                max_tokens=4096
            )
            
            analysis = response.content
            
            logger.info(
                "Files analyzed with images (Responses API)",
                file_count=len(files_content),
                images_count=len(images),
                analysis_length=len(analysis)
            )
            
            return analysis.strip()
            
        except Exception as e:
            logger.error(f"Failed to analyze files with images: {e}")
            # در صورت خطا، فقط محتوای خام را برگردان
            return self._fallback_analysis(files_content)
    
    def _build_analysis_prompt_for_mixed(
        self,
        text_files: List[Dict[str, Any]],
        images: List[Dict[str, Any]],
        user_query: str,
        language: str
    ) -> str:
        """ساخت prompt برای تحلیل فایل‌های ترکیبی (متن + تصویر)"""
        parts = []
        
        if language == "fa":
            parts.append(f"سوال کاربر: {user_query}\n")
            
            if text_files:
                parts.append(f"\n--- فایل‌های متنی ({len(text_files)} فایل) ---")
                for i, file_info in enumerate(text_files, 1):
                    parts.append(f"\nفایل {i}: {file_info['filename']}")
                    content = file_info.get('content', '')
                    if content:
                        max_length = 3000
                        if len(content) > max_length:
                            content = content[:max_length] + "\n... (ادامه دارد)"
                        parts.append(f"محتوا:\n{content}")
                    else:
                        parts.append("(محتوای متنی استخراج نشد)")
            
            if images:
                parts.append(f"\n--- تصاویر ({len(images)} تصویر) ---")
                for i, img in enumerate(images, 1):
                    parts.append(f"تصویر {i}: {img['filename']}")
                parts.append("\nلطفاً تصاویر بالا را تحلیل کن.")
            
            parts.append("\n\nلطفاً این فایل‌ها را تحلیل کن و اطلاعات مهم را استخراج کن.")
        else:
            parts.append(f"User's question: {user_query}\n")
            
            if text_files:
                parts.append(f"\n--- Text Files ({len(text_files)} files) ---")
                for i, file_info in enumerate(text_files, 1):
                    parts.append(f"\nFile {i}: {file_info['filename']}")
                    content = file_info.get('content', '')
                    if content:
                        max_length = 3000
                        if len(content) > max_length:
                            content = content[:max_length] + "\n... (continued)"
                        parts.append(f"Content:\n{content}")
                    else:
                        parts.append("(No text content extracted)")
            
            if images:
                parts.append(f"\n--- Images ({len(images)} images) ---")
                for i, img in enumerate(images, 1):
                    parts.append(f"Image {i}: {img['filename']}")
                parts.append("\nPlease analyze the images above.")
            
            parts.append("\n\nPlease analyze these files and extract important information.")
        
        return "\n".join(parts)
    
    async def analyze_image_with_vision(
        self,
        image_data: bytes,
        filename: str,
        user_query: str,
        language: str = "fa"
    ) -> str:
        """
        تحلیل تصویر با Vision API
        
        Args:
            image_data: داده باینری تصویر
            filename: نام فایل
            user_query: سوال کاربر
            language: زبان
            
        Returns:
            تحلیل تصویر
        """
        try:
            # تبدیل به base64
            image_base64 = base64.b64encode(image_data).decode('utf-8')
            
            # تهیه prompt
            vision_prompt = f"""تصویر ضمیمه شده را تحلیل کن.
سوال کاربر: {user_query}

لطفاً موارد زیر را ارائه کن:
1. توضیح کامل محتوای تصویر
2. متن موجود در تصویر (اگر وجود دارد)
3. ارتباط تصویر با سوال کاربر
4. نکات مهم و کلیدی

پاسخ را به زبان {'فارسی' if language == 'fa' else 'انگلیسی'} بده."""

            # ارسال به LLM با Vision
            messages = [
                Message(
                    role="user",
                    content=[
                        {
                            "type": "text",
                            "text": vision_prompt
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{image_base64}"
                            }
                        }
                    ]
                )
            ]
            
            # استفاده از Responses API برای تحلیل تصویر
            response = await self.llm.generate_responses_api(
                messages,
                reasoning_effort="medium"
            )
            analysis = response.content  # Extract content from LLMResponse
            
            logger.info(
                "Image analyzed with vision (Responses API)",
                filename=filename,
                analysis_length=len(analysis)
            )
            
            return analysis.strip()
            
        except Exception as e:
            logger.error(f"Failed to analyze image with vision: {e}")
            return f"خطا در تحلیل تصویر {filename}"
    
    def _get_system_prompt(self, language: str) -> str:
        """دریافت system prompt برای تحلیل فایل"""
        if language == "fa":
            return """تو یک دستیار هوشمند تحلیل اسناد هستی.
وظیفه‌ات تحلیل فایل‌های ضمیمه شده و استخراج اطلاعات مهم است.

برای هر فایل:
1. خلاصه محتوا را ارائه کن
2. نکات کلیدی و مهم را استخراج کن
3. ارتباط محتوا با سوال کاربر را مشخص کن
4. اگر جدول یا داده عددی وجود دارد، آن را تفسیر کن

پاسخ را مختصر، دقیق و کاربردی بنویس."""
        else:
            return """You are an intelligent document analysis assistant.
Your task is to analyze attached files and extract important information.

For each file:
1. Provide a summary of the content
2. Extract key points
3. Identify relevance to user's question
4. Interpret tables or numerical data if present

Keep the response concise, accurate and practical."""
    
    def _build_analysis_prompt(
        self,
        files_content: List[Dict[str, Any]],
        user_query: str,
        language: str
    ) -> str:
        """ساخت prompt برای تحلیل فایل‌ها"""
        parts = []
        
        if language == "fa":
            parts.append(f"سوال کاربر: {user_query}\n")
            parts.append(f"تعداد فایل‌های ضمیمه: {len(files_content)}\n")
            
            for i, file_info in enumerate(files_content, 1):
                parts.append(f"\n--- فایل {i}: {file_info['filename']} ---")
                parts.append(f"نوع: {file_info['file_type']}")
                
                if file_info.get('is_image'):
                    parts.append("(این فایل یک تصویر است)")
                
                content = file_info.get('content', '')
                if content:
                    # محدود کردن طول محتوا
                    max_length = 3000
                    if len(content) > max_length:
                        content = content[:max_length] + "\n... (ادامه دارد)"
                    parts.append(f"\nمحتوا:\n{content}")
                else:
                    parts.append("\n(محتوای متنی استخراج نشد)")
            
            parts.append("\n\nلطفاً این فایل‌ها را تحلیل کن و اطلاعات مهم را استخراج کن.")
        else:
            parts.append(f"User's question: {user_query}\n")
            parts.append(f"Number of attached files: {len(files_content)}\n")
            
            for i, file_info in enumerate(files_content, 1):
                parts.append(f"\n--- File {i}: {file_info['filename']} ---")
                parts.append(f"Type: {file_info['file_type']}")
                
                if file_info.get('is_image'):
                    parts.append("(This is an image file)")
                
                content = file_info.get('content', '')
                if content:
                    max_length = 3000
                    if len(content) > max_length:
                        content = content[:max_length] + "\n... (continued)"
                    parts.append(f"\nContent:\n{content}")
                else:
                    parts.append("\n(No text content extracted)")
            
            parts.append("\n\nPlease analyze these files and extract important information.")
        
        return "\n".join(parts)
    
    def _fallback_analysis(self, files_content: List[Dict[str, Any]]) -> str:
        """تحلیل ساده در صورت خطا در LLM"""
        parts = ["محتوای فایل‌های ضمیمه:\n"]
        
        for i, file_info in enumerate(files_content, 1):
            parts.append(f"\n{i}. {file_info['filename']} ({file_info['file_type']})")
            content = file_info.get('content', '')
            if content:
                # محدود کردن طول
                if len(content) > 500:
                    content = content[:500] + "..."
                parts.append(content)
        
        return "\n".join(parts)


# Singleton instance
_file_analysis_service = None


def get_file_analysis_service() -> FileAnalysisService:
    """Get file analysis service instance"""
    global _file_analysis_service
    if _file_analysis_service is None:
        _file_analysis_service = FileAnalysisService()
    return _file_analysis_service
